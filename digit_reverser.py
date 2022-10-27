from docx import Document

def docx_define(path=None):
    return Document(path)

def digits_reverser(doc):
    # new docx
    document = docx_define()
    for p in doc.paragraphs:
        txtchain = ''
        number_checker = 0
        num_stack = []
        for text in p.text:
            if text in "0123456789۱۲۳۴۵۶۷۸۹۰":
                if number_checker == 0:
                    num_stack.append(text)
            else:
                if len(num_stack)>0:
                    number = "".join(str(i) for i in num_stack[::-1])
                    txtchain += number
                    num_stack = []
                else:
                    txtchain = txtchain + text

        document.add_paragraph(txtchain)
    document.save("new1.docx")